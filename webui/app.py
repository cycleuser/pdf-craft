# -*- coding: utf-8 -*-
import os
import uuid
import json
import threading
import logging
import pickle
import time
import multiprocessing as mp
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_from_directory, redirect, url_for
from werkzeug.utils import secure_filename
from pdf_craft import PDFPageExtractor, MarkDownWriter
from optimization_config import OptimizationConfigManager, auto_configure_optimization, get_system_performance_report, get_preset_config, list_preset_configs

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
app.config['RESULTS_FOLDER'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'results')
app.config['MODEL_DIR'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'models')
app.config['JOBS_FILE'] = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'jobs.pkl')
app.config['ALLOWED_EXTENSIONS'] = {'pdf'}
app.config['USE_GPU'] = False  # 默认使用CPU，可以在这里修改为True来使用GPU
app.config['BATCH_SIZE'] = 5  # 批处理大小，可以根据系统性能调整

# 性能优化配置
app.config['MAX_WORKERS'] = min(32, (os.cpu_count() or 1) + 4)  # 最大工作线程数
app.config['GPU_BATCH_SIZE'] = 8  # GPU批处理大小
app.config['CPU_BATCH_SIZE'] = 4  # CPU批处理大小
app.config['MEMORY_LIMIT_GB'] = 8  # 内存限制（GB）
app.config['ENABLE_MULTIPROCESSING'] = True  # 启用多进程处理
app.config['PROCESS_POOL_SIZE'] = min(4, os.cpu_count() or 1)  # 进程池大小
app.config['PRELOAD_MODELS'] = True  # 预加载模型
app.config['ENABLE_MIXED_PRECISION'] = True  # 启用混合精度
app.config['OPTIMIZE_MEMORY'] = True  # 启用内存优化

# OCR级别配置
OCR_LEVELS = {
    'fast': {'name': '快速', 'description': '快速OCR处理，适合简单文档'},
    'standard': {'name': '标准 (默认)', 'description': '标准OCR处理，平衡速度和准确性'},
    'accurate': {'name': '精确', 'description': '高精度OCR处理，适合复杂文档'},
    'detailed': {'name': '详细', 'description': '最详细的OCR处理，包含更多元数据'},
}

# 表格提取格式
TABLE_FORMATS = {
    'none': {'name': '不提取', 'description': '跳过表格提取'},
    'simple': {'name': '简单', 'description': '基本表格提取'},
    'standard': {'name': '标准 (默认)', 'description': '标准表格处理，保持格式'},
    'advanced': {'name': '高级', 'description': '高级表格处理，包含样式信息'},
}

# Ollama模型配置
OLLAMA_MODELS = {
    'none': {'name': '不使用LLM', 'description': '不使用大语言模型进行处理'},
    'llama3': {'name': 'Llama 3', 'description': '使用Llama 3模型进行处理'},
    'mistral': {'name': 'Mistral', 'description': '使用Mistral模型进行处理'},
    'gemma': {'name': 'Gemma', 'description': '使用Gemma模型进行处理'},
    'phi3': {'name': 'Phi-3', 'description': '使用Phi-3模型进行处理'},
}

# 默认配置
DEFAULT_CONFIG = {
    'ocr_level': 'standard',
    'extract_formula': False,
    'extract_table_format': 'standard',
    'ollama_model': 'none',
}

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)
os.makedirs(app.config['MODEL_DIR'], exist_ok=True)

# Store job status
jobs = {}
# 处理队列
job_queue = []
# 当前正在处理的作业数量
active_jobs = 0
# 线程锁
queue_lock = threading.Lock()
# 当前模型配置
current_config = DEFAULT_CONFIG.copy()

# 全局模型缓存
model_cache = {}
model_cache_lock = threading.Lock()

# 性能监控
performance_stats = {
    'total_pages_processed': 0,
    'total_processing_time': 0,
    'average_pages_per_second': 0,
    'gpu_utilization': 0,
    'memory_usage': 0
}

# 从文件加载作业状态
def load_jobs():
    global jobs, job_queue
    if os.path.exists(app.config['JOBS_FILE']):
        try:
            with open(app.config['JOBS_FILE'], 'rb') as f:
                jobs = pickle.load(f)
            logger.info(f"已从文件加载 {len(jobs)} 个作业状态")

            # 检查文件是否存在
            for job_id, job in list(jobs.items()):
                if 'file_path' in job and not os.path.exists(job['file_path']):
                    logger.warning(f"作业 {job_id} 的文件不存在，移除该作业")
                    del jobs[job_id]

                if job['status'] == 'processing':
                    logger.warning(f"作业 {job_id} 处于处理中状态，重置为排队状态")
                    job['status'] = 'queued'
                    job['progress'] = {
                        'current_page': 0,
                        'total_pages': 0,
                        'percentage': 0
                    }
                    # 将作业重新加入队列
                    job_queue.append(job_id)

                if job['status'] == 'queued':
                    # 将排队中的作业加入队列
                    job_queue.append(job_id)
        except Exception as e:
            logger.error(f"加载作业状态失败: {str(e)}")
            jobs = {}
            job_queue = []

# 保存作业状态到文件
def save_jobs():
    try:
        with open(app.config['JOBS_FILE'], 'wb') as f:
            pickle.dump(jobs, f)
        logger.info(f"已保存 {len(jobs)} 个作业状态到文件")
    except Exception as e:
        logger.error(f"保存作业状态失败: {str(e)}")

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

class OptimizedProgressReporter:
    def __init__(self, job_id):
        self.job_id = job_id
        self.current_page = 0
        self.total_pages = 0
        self.last_save_time = 0
        self.start_time = time.time()
        self.page_times = []

    def report(self, current_page, total_pages):
        self.current_page = current_page
        self.total_pages = total_pages
        
        # 计算处理速度
        current_time = time.time()
        if len(self.page_times) > 0:
            elapsed = current_time - self.page_times[-1]
            self.page_times.append(current_time)
            if len(self.page_times) > 10:  # 保持最近10页的时间记录
                self.page_times.pop(0)
        else:
            self.page_times.append(current_time)
        
        # 计算平均速度
        if len(self.page_times) > 1:
            total_time = self.page_times[-1] - self.page_times[0]
            pages_per_second = (len(self.page_times) - 1) / total_time if total_time > 0 else 0
        else:
            pages_per_second = 0
        
        jobs[self.job_id]['progress'] = {
            'current_page': current_page,
            'total_pages': total_pages,
            'percentage': int((current_page / total_pages) * 100) if total_pages > 0 else 0,
            'pages_per_second': round(pages_per_second, 2),
            'estimated_remaining_time': int((total_pages - current_page) / pages_per_second) if pages_per_second > 0 else 0
        }
        
        logger.info(f"Job {self.job_id}: Processing page {current_page}/{total_pages} ({pages_per_second:.2f} pages/sec)")

        # 每5秒保存一次状态，而不是每次更新都保存
        if current_time - self.last_save_time > 5:
            save_jobs()
            self.last_save_time = current_time

def get_optimal_device_config():
    """获取最优设备配置"""
    # 首先检查是否启用了GPU
    if not app.config.get('USE_GPU', False):
        # 如果禁用GPU，直接返回CPU配置
        cpu_count = os.cpu_count() or 1
        return {
            'device': 'cpu',
            'batch_size': app.config.get('CPU_BATCH_SIZE', min(4, cpu_count)),
            'cpu_count': cpu_count,
            'mixed_precision': False
        }
    
    try:
        import torch
        if torch.cuda.is_available():
            gpu_count = torch.cuda.device_count()
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3  # GB
            
            # 根据GPU内存调整批处理大小，但优先使用配置的值
            configured_batch_size = app.config.get('GPU_BATCH_SIZE')
            if configured_batch_size:
                batch_size = configured_batch_size
            else:
                # 自动根据GPU内存调整
                if gpu_memory >= 16:
                    batch_size = 16
                elif gpu_memory >= 8:
                    batch_size = 8
                elif gpu_memory >= 4:
                    batch_size = 4
                else:
                    batch_size = 2
                
            return {
                'device': 'cuda',
                'batch_size': batch_size,
                'gpu_count': gpu_count,
                'gpu_memory': gpu_memory,
                'mixed_precision': app.config.get('ENABLE_MIXED_PRECISION', False)
            }
    except ImportError:
        logger.warning("PyTorch未安装，无法使用GPU")
    except Exception as e:
        logger.warning(f"GPU检测失败: {str(e)}")
    
    # CPU配置
    cpu_count = os.cpu_count() or 1
    return {
        'device': 'cpu',
        'batch_size': app.config.get('CPU_BATCH_SIZE', min(4, cpu_count)),
        'cpu_count': cpu_count,
        'mixed_precision': False
    }

def create_optimized_extractor(device_config, model_dir):
    """创建优化的PDF提取器"""
    with model_cache_lock:
        cache_key = f"{device_config['device']}"
        
        if cache_key in model_cache and app.config['PRELOAD_MODELS']:
            logger.info(f"使用缓存的模型: {cache_key}")
            return model_cache[cache_key]
        
        try:
            # 检查PDFPageExtractor支持的参数
            import inspect
            extractor_signature = inspect.signature(PDFPageExtractor.__init__)
            supported_params = list(extractor_signature.parameters.keys())
            
            # 构建参数字典，只包含支持的参数
            extractor_params = {}
            
            # 设备参数
            if 'device' in supported_params:
                extractor_params['device'] = device_config['device']
            
            # 模型目录参数
            if 'model_dir_path' in supported_params:
                extractor_params['model_dir_path'] = model_dir
            elif 'model_dir' in supported_params:
                extractor_params['model_dir'] = model_dir
            
            # 批处理大小参数（如果支持）
            if 'batch_size' in supported_params:
                extractor_params['batch_size'] = device_config['batch_size']
            
            logger.info(f"创建提取器，支持的参数: {supported_params}")
            logger.info(f"使用的参数: {extractor_params}")
            
            # 创建提取器
            extractor = PDFPageExtractor(**extractor_params)
            
            # 如果支持，启用优化选项
            if hasattr(extractor, 'enable_optimization'):
                try:
                    extractor.enable_optimization(
                        mixed_precision=device_config.get('mixed_precision', False),
                        memory_efficient=app.config['OPTIMIZE_MEMORY'],
                        parallel_processing=True
                    )
                    logger.info("已启用提取器优化选项")
                except Exception as opt_e:
                    logger.warning(f"启用优化选项失败: {str(opt_e)}")
            
            # 设置批处理大小（如果支持）
            if hasattr(extractor, 'set_batch_size'):
                try:
                    extractor.set_batch_size(device_config['batch_size'])
                    logger.info(f"已设置批处理大小: {device_config['batch_size']}")
                except Exception as batch_e:
                    logger.warning(f"设置批处理大小失败: {str(batch_e)}")
            
            if app.config['PRELOAD_MODELS']:
                model_cache[cache_key] = extractor
                logger.info(f"模型已缓存: {cache_key}")
            
            return extractor
            
        except Exception as e:
            logger.error(f"创建优化提取器失败: {str(e)}")
            # 回退到最基本的配置
            try:
                return PDFPageExtractor(model_dir_path=model_dir)
            except Exception as fallback_e:
                logger.error(f"回退创建提取器也失败: {str(fallback_e)}")
                # 最后的回退，不指定任何参数
                return PDFPageExtractor()

def process_pdf_pages_parallel(extractor, pdf_path, progress_reporter, job_config):
    """并行处理PDF页面"""
    try:
        # 获取PDF页面数
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        doc.close()
        
        logger.info(f"PDF页面总数: {total_pages}")
        
        # 检查是否支持单页处理
        supports_page_extraction = hasattr(extractor, 'extract_page')
        
        # 检查当前pdf_craft库是否支持并行处理
        if not supports_page_extraction:
            logger.warning("当前pdf_craft版本不支持单页提取，使用标准串行处理")
            return list(extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report))
            
        # 对于小文档，直接使用标准处理
        if total_pages <= 10:
            logger.info("小文档，使用标准处理")
            return list(extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report))
        
        # 尝试使用优化的处理方式
        logger.info("尝试使用优化的处理方式")
        return list(extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report))
        
    except Exception as e:
        logger.error(f"并行处理失败，回退到串行处理: {str(e)}")
        # 回退到原始处理方式
        try:
            return list(extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report))
        except Exception as fallback_e:
            logger.error(f"串行处理也失败: {str(fallback_e)}")
            return []

def process_pdf_batch(pdf_path, start_page, end_page, device_config, model_dir, job_config):
    """处理PDF批次（多进程函数）"""
    try:
        # 在子进程中创建提取器
        extractor = create_optimized_extractor(device_config, model_dir)
        
        # 处理指定页面范围
        results = []
        for page_num in range(start_page, end_page):
            try:
                page_result = process_single_page_with_extractor(extractor, pdf_path, page_num, job_config)
                results.append(page_result)
            except Exception as e:
                logger.error(f"批次中页面 {page_num} 处理失败: {str(e)}")
        
        return results
        
    except Exception as e:
        logger.error(f"批次处理失败: {str(e)}")
        return []

def process_single_page(pdf_path, page_num, device_config, model_dir, job_config):
    """处理单个页面"""
    try:
        extractor = create_optimized_extractor(device_config, model_dir)
        return process_single_page_with_extractor(extractor, pdf_path, page_num, job_config)
    except Exception as e:
        logger.error(f"单页处理失败: {str(e)}")
        return None

def process_single_page_with_extractor(extractor, pdf_path, page_num, job_config):
    """使用指定提取器处理单个页面"""
    try:
        # 这里需要根据pdf_craft库的实际API调整
        # 假设支持单页提取
        if hasattr(extractor, 'extract_page'):
            return extractor.extract_page(pdf_path, page_num)
        else:
            # 如果不支持单页提取，使用原始方法
            return None
    except Exception as e:
        logger.error(f"页面提取失败: {str(e)}")
        return None

# 初始化Ollama LLM
def init_ollama_llm(model_name):
    """初始化Ollama LLM"""
    if model_name == 'none':
        logger.info("不使用LLM模型")
        return None

    try:
        # 尝试导入LLM，如果不存在则返回None
        try:
            from pdf_craft import LLM
            logger.info(f"成功导入LLM模块")

            # 初始化Ollama LLM
            llm = LLM(
                url="http://localhost:11434",  # Ollama默认地址
                model=model_name,
                token_encoding="o200k_base"
            )
            logger.info(f"成功初始化Ollama LLM模型: {model_name}")
            return llm
        except ImportError:
            logger.warning("pdf_craft库中不存在LLM模块，跳过LLM初始化")
            return None
        except Exception as e:
            logger.error(f"初始化Ollama LLM失败: {str(e)}")
            return None
    except Exception as e:
        logger.error(f"初始化Ollama LLM时出错: {str(e)}")
        return None

# 检查Ollama是否可用
def check_ollama_available():
    """检查Ollama服务是否可用"""
    import requests
    try:
        response = requests.get("http://localhost:11434/api/version", timeout=2)
        if response.status_code == 200:
            # 获取可用模型
            models_response = requests.get("http://localhost:11434/api/tags", timeout=2)
            if models_response.status_code == 200:
                models_data = models_response.json()
                available_models = [model['name'] for model in models_data.get('models', [])]
                logger.info(f"Ollama可用，已发现模型: {available_models}")
                return True, available_models
            return True, []
        return False, []
    except:
        return False, []

def process_pdf(job_id, pdf_path, output_dir):
    global active_jobs, performance_stats
    start_time = time.time()
    
    try:
        # Update job status
        jobs[job_id]['status'] = 'processing'
        save_jobs()  # 保存状态变更

        # Create a unique folder for this job's results
        job_result_dir = os.path.join(app.config['RESULTS_FOLDER'], job_id)
        os.makedirs(job_result_dir, exist_ok=True)

        # Get the filename without extension
        pdf_filename = os.path.basename(pdf_path)
        # 使用原始文件名（如果存在）
        if 'original_filename' in jobs[job_id]:
            original_filename = jobs[job_id]['original_filename']
            base_filename = os.path.splitext(original_filename)[0]
        else:
            base_filename = os.path.splitext(pdf_filename)[0]

        # Create markdown output path
        markdown_path = os.path.join(job_result_dir, f"{base_filename}.md")
        images_dir = "images"

        # Create optimized progress reporter
        progress_reporter = OptimizedProgressReporter(job_id)

        # 获取最优设备配置
        device_config = get_optimal_device_config()
        device = device_config['device']
        
        logger.info(f"Job {job_id}: Using optimized config: {device_config}")

        # 获取作业配置
        job_config = jobs[job_id].get('config', current_config)

        # 获取Ollama模型配置
        ollama_model = job_config.get('ollama_model', 'none')

        # 初始化LLM（如果选择了模型）
        llm = None
        if ollama_model != 'none':
            llm = init_ollama_llm(ollama_model)

        # 创建优化的PDF提取器
        extractor = create_optimized_extractor(device_config, app.config['MODEL_DIR'])

        # 记录使用的配置
        config_info = {
            'device': device,
            'batch_size': device_config['batch_size'],
            'ollama_model': ollama_model,
            'optimization_enabled': True,
            'mixed_precision': device_config.get('mixed_precision', False)
        }

        # 获取并记录模型信息
        model_files = os.listdir(app.config['MODEL_DIR']) if os.path.exists(app.config['MODEL_DIR']) else []
        logger.info(f"Job {job_id}: Models available: {model_files}")

        # 记录开始处理PDF
        logger.info(f"Job {job_id}: Starting optimized PDF processing with config: {config_info}")

        # 使用稳定的处理方式
        try:
            logger.info(f"开始处理PDF文件: {pdf_path}")
            with MarkDownWriter(markdown_path, images_dir, "utf-8") as md:
                # 根据配置选择处理方式
                if app.config['ENABLE_MULTIPROCESSING']:
                    # 尝试使用优化的处理方式
                    logger.info("使用优化的处理方式")
                    blocks = process_pdf_pages_parallel(extractor, pdf_path, progress_reporter, job_config)
                    for block in blocks:
                        if block:  # 确保block不为None
                            md.write(block)
                else:
                    # 使用标准串行处理
                    logger.info("使用标准串行处理")
                    for block in extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report):
                        md.write(block)
                        
            logger.info(f"PDF处理完成: {markdown_path}")
        except Exception as e:
            logger.error(f"PDF处理失败，尝试最基本的处理方式: {str(e)}")
            try:
                # 最后的回退方案：使用最基本的处理方式
                with MarkDownWriter(markdown_path, images_dir, "utf-8") as md:
                    for block in extractor.extract(pdf=pdf_path, report_progress=progress_reporter.report):
                        md.write(block)
                logger.info("使用基本处理方式成功")
            except Exception as fallback_e:
                logger.error(f"基本处理方式也失败: {str(fallback_e)}")
                raise fallback_e

        # 生成Word文档
        word_path = os.path.join(job_result_dir, f"{base_filename}.docx")
        create_word_from_markdown(markdown_path, word_path, os.path.join(job_result_dir, images_dir))

        # 生成带文本的PDF
        text_pdf_path = os.path.join(job_result_dir, f"{base_filename}_text.pdf")
        create_text_pdf_from_markdown(markdown_path, text_pdf_path, os.path.join(job_result_dir, images_dir))

        # 生成完整压缩包
        complete_zip_path = os.path.join(job_result_dir, f"{base_filename}_complete.zip")
        create_complete_zip(job_result_dir, complete_zip_path, base_filename)

        # 计算处理时间和性能统计
        end_time = time.time()
        processing_time = end_time - start_time
        
        # 更新性能统计
        total_pages = progress_reporter.total_pages
        if total_pages > 0:
            performance_stats['total_pages_processed'] += total_pages
            performance_stats['total_processing_time'] += processing_time
            performance_stats['average_pages_per_second'] = performance_stats['total_pages_processed'] / performance_stats['total_processing_time']

        # Update job status to completed
        jobs[job_id]['status'] = 'completed'
        jobs[job_id]['result'] = {
            'markdown_path': markdown_path,
            'word_path': word_path,
            'text_pdf_path': text_pdf_path,
            'complete_zip_path': complete_zip_path,
            'job_result_dir': job_result_dir,
            'filename': f"{base_filename}.md",
            'word_filename': f"{base_filename}.docx",
            'text_pdf_filename': f"{base_filename}_text.pdf",
            'complete_zip_filename': f"{base_filename}_complete.zip",
            'device_used': device,
            'model_dir': app.config['MODEL_DIR'],
            'processing_time': processing_time,
            'config': config_info,
            'performance': {
                'total_pages': total_pages,
                'pages_per_second': total_pages / processing_time if processing_time > 0 else 0,
                'optimization_used': True
            }
        }
        logger.info(f"Job {job_id}: Completed successfully in {processing_time:.2f} seconds ({total_pages / processing_time if processing_time > 0 else 0:.2f} pages/sec)")
        save_jobs()  # 保存完成状态

    except Exception as e:
        # Update job status to failed
        jobs[job_id]['status'] = 'failed'
        jobs[job_id]['error'] = str(e)
        logger.error(f"Job {job_id}: Failed with error: {str(e)}")
        save_jobs()  # 保存失败状态
    finally:
        # 减少活跃作业计数
        with queue_lock:
            active_jobs -= 1
        # 处理队列中的下一个作业
        process_next_job()

def create_word_from_markdown(markdown_path, word_path, images_dir):
    """将Markdown转换为Word文档"""
    try:
        from markdown import markdown
        from docx import Document
        from docx.shared import Inches
        import re

        # 读取Markdown文件
        with open(markdown_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 创建Word文档
        doc = Document()

        # 处理图片引用
        image_pattern = r'!\[(.*?)\]\((.*?)\)'

        # 分段处理Markdown内容
        paragraphs = md_content.split('\n\n')

        for para in paragraphs:
            if not para.strip():
                continue

            # 处理图片
            img_match = re.search(image_pattern, para)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)

                # 获取图片的完整路径
                if img_path.startswith('images/'):
                    full_img_path = os.path.join(images_dir, os.path.basename(img_path))
                    if os.path.exists(full_img_path):
                        doc.add_picture(full_img_path, width=Inches(6))
                        if alt_text:
                            doc.add_paragraph(alt_text)
                continue

            # 处理标题
            if para.startswith('# '):
                doc.add_heading(para[2:], level=1)
            elif para.startswith('## '):
                doc.add_heading(para[3:], level=2)
            elif para.startswith('### '):
                doc.add_heading(para[4:], level=3)
            else:
                # 处理普通段落
                doc.add_paragraph(para)

        # 保存Word文档
        doc.save(word_path)
        logger.info(f"Word文档已保存: {word_path}")
        return True
    except Exception as e:
        logger.error(f"创建Word文档失败: {str(e)}")
        return False

def create_text_pdf_from_markdown(markdown_path, pdf_path, images_dir):
    """将Markdown转换为带文本的PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        import re

        # 读取Markdown文件
        with open(markdown_path, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # 创建PDF文档
        doc = SimpleDocTemplate(pdf_path, pagesize=A4)
        styles = getSampleStyleSheet()

        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=8
        )

        story = []

        # 处理图片引用
        image_pattern = r'!\[(.*?)\]\((.*?)\)'

        # 分段处理Markdown内容
        paragraphs = md_content.split('\n\n')

        for para in paragraphs:
            if not para.strip():
                continue

            # 处理图片
            img_match = re.search(image_pattern, para)
            if img_match:
                alt_text = img_match.group(1)
                img_path = img_match.group(2)

                # 获取图片的完整路径
                if img_path.startswith('images/'):
                    full_img_path = os.path.join(images_dir, os.path.basename(img_path))
                    if os.path.exists(full_img_path):
                        try:
                            img = Image(full_img_path, width=6*inch, height=4*inch)
                            story.append(img)
                            if alt_text:
                                story.append(Paragraph(alt_text, styles['Caption']))
                            story.append(Spacer(1, 12))
                        except Exception as e:
                            logger.warning(f"无法添加图片到PDF: {full_img_path}, 错误: {str(e)}")
                continue

            # 处理标题
            if para.startswith('# '):
                story.append(Paragraph(para[2:], title_style))
            elif para.startswith('## '):
                story.append(Paragraph(para[3:], heading_style))
            elif para.startswith('### '):
                story.append(Paragraph(para[4:], styles['Heading3']))
            else:
                # 处理普通段落
                # 简单处理Markdown格式
                para = para.replace('**', '<b>').replace('**', '</b>')
                para = para.replace('*', '<i>').replace('*', '</i>')
                story.append(Paragraph(para, styles['Normal']))
                story.append(Spacer(1, 6))

        # 构建PDF
        doc.build(story)
        logger.info(f"带文本PDF已保存: {pdf_path}")
        return True
    except Exception as e:
        logger.error(f"创建带文本PDF失败: {str(e)}")
        return False

def create_complete_zip(job_result_dir, zip_path, base_filename):
    """创建包含所有文件的完整压缩包"""
    try:
        import zipfile

        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            # 添加所有文件到压缩包
            for root, dirs, files in os.walk(job_result_dir):
                for file in files:
                    if file.endswith('.zip'):  # 跳过其他zip文件
                        continue
                    file_path = os.path.join(root, file)
                    # 计算相对路径
                    arcname = os.path.relpath(file_path, job_result_dir)
                    zipf.write(file_path, arcname)

        logger.info(f"完整压缩包已创建: {zip_path}")
        return True
    except Exception as e:
        logger.error(f"创建完整压缩包失败: {str(e)}")
        return False

def process_next_job():
    """处理队列中的下一个作业"""
    global active_jobs
    with queue_lock:
        # 检查是否有排队的作业
        if job_queue and active_jobs < app.config['BATCH_SIZE']:
            # 获取下一个作业ID
            job_id = job_queue.pop(0)
            if job_id in jobs:
                # 增加活跃作业计数
                active_jobs += 1
                # 启动处理线程
                thread = threading.Thread(
                    target=process_pdf,
                    args=(job_id, jobs[job_id]['file_path'], app.config['RESULTS_FOLDER'])
                )
                thread.start()
                logger.info(f"开始处理作业 {job_id}, 当前活跃作业数: {active_jobs}")
            else:
                logger.warning(f"作业 {job_id} 不存在，从队列中移除")

@app.route('/')
def index():
    # 检查Ollama是否可用
    ollama_available, available_models = check_ollama_available()

    # 如果Ollama可用，更新模型列表
    if ollama_available and available_models:
        # 更新OLLAMA_MODELS，保留'none'选项
        none_option = OLLAMA_MODELS['none']
        OLLAMA_MODELS.clear()
        OLLAMA_MODELS['none'] = none_option

        # 添加可用模型
        for model in available_models:
            model_id = model.lower().replace(':', '-')
            OLLAMA_MODELS[model_id] = {
                'name': model,
                'description': f'使用{model}模型进行处理'
            }

    # 传递设备信息和模型配置到模板
    return render_template('index.html',
                          use_gpu=app.config['USE_GPU'],
                          ocr_levels=OCR_LEVELS,
                          table_formats=TABLE_FORMATS,
                          ollama_models=OLLAMA_MODELS,
                          ollama_available=ollama_available,
                          current_config=current_config)

@app.route('/toggle_gpu', methods=['POST'])
def toggle_gpu():
    app.config['USE_GPU'] = not app.config['USE_GPU']
    logger.info(f"GPU acceleration {'enabled' if app.config['USE_GPU'] else 'disabled'}")
    return jsonify({'use_gpu': app.config['USE_GPU']})

@app.route('/update_config', methods=['POST'])
def update_config():
    """更新模型配置"""
    global current_config
    try:
        data = request.get_json()

        # 验证OCR级别
        if 'ocr_level' in data and data['ocr_level'] in OCR_LEVELS:
            current_config['ocr_level'] = data['ocr_level']

        # 验证表格提取格式
        if 'extract_table_format' in data and data['extract_table_format'] in TABLE_FORMATS:
            current_config['extract_table_format'] = data['extract_table_format']

        # 验证是否提取公式
        if 'extract_formula' in data:
            current_config['extract_formula'] = bool(data['extract_formula'])

        # 验证Ollama模型
        if 'ollama_model' in data and data['ollama_model'] in OLLAMA_MODELS:
            current_config['ollama_model'] = data['ollama_model']

        logger.info(f"模型配置已更新: {current_config}")

        return jsonify({
            'success': True,
            'config': current_config
        })
    except Exception as e:
        logger.error(f"更新模型配置失败: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/reset_config', methods=['POST'])
def reset_config():
    """重置模型配置为默认值"""
    global current_config
    try:
        current_config = DEFAULT_CONFIG.copy()
        logger.info(f"模型配置已重置为默认值: {current_config}")

        return jsonify({
            'success': True,
            'config': current_config
        })
    except Exception as e:
        logger.error(f"重置模型配置失败: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/system_info')
def system_info():
    # 检查CUDA是否可用
    try:
        import torch
        cuda_available = torch.cuda.is_available()
        cuda_device_count = torch.cuda.device_count() if cuda_available else 0
        cuda_device_name = torch.cuda.get_device_name(0) if cuda_available and cuda_device_count > 0 else "N/A"
        
        # 获取GPU内存信息
        gpu_memory_info = {}
        if cuda_available:
            for i in range(cuda_device_count):
                props = torch.cuda.get_device_properties(i)
                gpu_memory_info[f'gpu_{i}'] = {
                    'name': props.name,
                    'total_memory': props.total_memory / 1024**3,  # GB
                    'allocated_memory': torch.cuda.memory_allocated(i) / 1024**3 if torch.cuda.is_initialized() else 0,
                    'cached_memory': torch.cuda.memory_reserved(i) / 1024**3 if torch.cuda.is_initialized() else 0
                }
    except ImportError:
        cuda_available = False
        cuda_device_count = 0
        cuda_device_name = "PyTorch not installed"
        gpu_memory_info = {}

    # 获取系统信息
    import platform
    import sys
    import psutil
    
    os_info = f"{platform.system()} {platform.release()}"
    python_version = f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}"
    
    # 获取CPU和内存信息
    cpu_info = {
        'cpu_count': os.cpu_count(),
        'cpu_percent': psutil.cpu_percent(interval=1),
        'memory_total': psutil.virtual_memory().total / 1024**3,  # GB
        'memory_available': psutil.virtual_memory().available / 1024**3,  # GB
        'memory_percent': psutil.virtual_memory().percent
    }

    # 获取模型目录信息
    model_dir = app.config['MODEL_DIR']
    model_files = os.listdir(model_dir) if os.path.exists(model_dir) else []

    # 获取队列信息
    queue_info = {
        'queued_jobs': len(job_queue),
        'active_jobs': active_jobs,
        'batch_size': app.config['BATCH_SIZE'],
        'max_workers': app.config['MAX_WORKERS'],
        'process_pool_size': app.config['PROCESS_POOL_SIZE']
    }

    # 检查Ollama是否可用
    ollama_available, available_models = check_ollama_available()
    
    # 获取优化配置
    device_config = get_optimal_device_config()
    
    # 获取性能统计
    performance_info = performance_stats.copy()

    return jsonify({
        'use_gpu': app.config['USE_GPU'],
        'cuda_available': cuda_available,
        'cuda_device_count': cuda_device_count,
        'cuda_device_name': cuda_device_name,
        'gpu_memory_info': gpu_memory_info,
        'cpu_info': cpu_info,
        'os_info': os_info,
        'python_version': python_version,
        'model_dir': model_dir,
        'model_files': model_files,
        'queue_info': queue_info,
        'current_config': current_config,
        'ollama_available': ollama_available,
        'ollama_models': available_models if ollama_available else [],
        'optimization_config': {
            'multiprocessing_enabled': app.config['ENABLE_MULTIPROCESSING'],
            'mixed_precision_enabled': app.config['ENABLE_MIXED_PRECISION'],
            'memory_optimization_enabled': app.config['OPTIMIZE_MEMORY'],
            'model_preloading_enabled': app.config['PRELOAD_MODELS'],
            'optimal_device_config': device_config
        },
        'performance_stats': performance_info,
        'model_cache_info': {
            'cached_models': list(model_cache.keys()),
            'cache_size': len(model_cache)
        }
    })

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'files[]' not in request.files:
        return jsonify({'error': 'No files part'}), 400

    files = request.files.getlist('files[]')

    if not files or files[0].filename == '':
        return jsonify({'error': 'No files selected'}), 400

    # 获取作业配置
    job_config = {}
    if request.form.get('use_custom_config') == 'true':
        # 使用自定义配置
        job_config['ocr_level'] = request.form.get('ocr_level', current_config['ocr_level'])
        job_config['extract_table_format'] = request.form.get('extract_table_format', current_config['extract_table_format'])
        job_config['extract_formula'] = request.form.get('extract_formula') == 'true'
        job_config['ollama_model'] = request.form.get('ollama_model', current_config['ollama_model'])
    else:
        # 使用全局配置
        job_config = current_config.copy()

    uploaded_files = []
    for file in files:
        if file and allowed_file(file.filename):
            # Generate a unique ID for the job
            job_id = str(uuid.uuid4())

            # 保留原始文件名，包括中文
            original_filename = file.filename
            # 生成安全的文件名用于存储
            safe_filename = secure_filename(file.filename)
            # 如果安全文件名为空或只有扩展名，使用UUID
            if not safe_filename or safe_filename.startswith('.'):
                file_ext = os.path.splitext(original_filename)[1]
                safe_filename = f"{job_id}{file_ext}"
            
            # 确保文件名唯一
            filename = safe_filename

            # Save the file
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(file_path)
            logger.info(f"File uploaded: {filename}, Job ID: {job_id}")

            # Create job entry
            jobs[job_id] = {
                'id': job_id,
                'filename': filename,
                'original_filename': original_filename,  # 保存原始文件名
                'file_path': file_path,
                'status': 'queued',
                'progress': {
                    'current_page': 0,
                    'total_pages': 0,
                    'percentage': 0
                },
                'created_at': time.time(),
                'config': job_config
            }

            # 将作业添加到队列
            with queue_lock:
                job_queue.append(job_id)

            # 保存作业状态
            save_jobs()

            uploaded_files.append({
                'job_id': job_id,
                'filename': filename,
                'original_filename': original_filename,
                'config': job_config
            })

    # 启动队列处理
    process_next_job()

    return jsonify({'success': True, 'message': 'Files uploaded successfully', 'jobs': uploaded_files})

@app.route('/jobs')
def get_jobs():
    return jsonify({'jobs': list(jobs.values())})

@app.route('/clear_all_jobs', methods=['POST'])
def clear_all_jobs():
    """清空所有任务和相关文件"""
    global jobs, job_queue, active_jobs
    try:
        with queue_lock:
            # 清空任务队列和状态
            jobs.clear()
            job_queue.clear()
            active_jobs = 0

            # 删除任务状态文件
            if os.path.exists(app.config['JOBS_FILE']):
                os.remove(app.config['JOBS_FILE'])

            # 清空上传文件夹
            upload_folder = app.config['UPLOAD_FOLDER']
            if os.path.exists(upload_folder):
                for filename in os.listdir(upload_folder):
                    file_path = os.path.join(upload_folder, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                        elif os.path.isdir(file_path):
                            import shutil
                            shutil.rmtree(file_path)
                    except Exception as e:
                        logger.warning(f"删除上传文件失败: {file_path}, 错误: {str(e)}")

            # 清空结果文件夹
            results_folder = app.config['RESULTS_FOLDER']
            if os.path.exists(results_folder):
                for filename in os.listdir(results_folder):
                    file_path = os.path.join(results_folder, filename)
                    try:
                        if os.path.isfile(file_path):
                            os.remove(file_path)
                        elif os.path.isdir(file_path):
                            import shutil
                            shutil.rmtree(file_path)
                    except Exception as e:
                        logger.warning(f"删除结果文件失败: {file_path}, 错误: {str(e)}")

        logger.info("所有任务和文件已清空")
        return jsonify({'success': True, 'message': '所有任务和文件已清空'})
    except Exception as e:
        logger.error(f"清空任务失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/job/<job_id>')
def get_job(job_id):
    if job_id in jobs:
        return jsonify(jobs[job_id])
    return jsonify({'error': 'Job not found'}), 404

@app.route('/job/<job_id>', methods=['DELETE'])
def delete_job(job_id):
    if job_id not in jobs:
        return jsonify({'success': False, 'error': 'Job not found'}), 404

    try:
        # 从队列中移除作业（如果存在）
        with queue_lock:
            if job_id in job_queue:
                job_queue.remove(job_id)

        # 删除作业文件
        job = jobs[job_id]
        if 'file_path' in job and os.path.exists(job['file_path']):
            os.remove(job['file_path'])

        # 删除结果文件夹
        if 'result' in job and 'job_result_dir' in job['result']:
            import shutil
            result_dir = job['result']['job_result_dir']
            if os.path.exists(result_dir):
                shutil.rmtree(result_dir)

        # 从作业字典中删除
        del jobs[job_id]

        # 保存状态
        save_jobs()

        logger.info(f"作业 {job_id} 已删除")
        return jsonify({'success': True, 'message': 'Job deleted successfully'})

    except Exception as e:
        logger.error(f"删除作业 {job_id} 失败: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/download/<job_id>')
def download_file(job_id):
    if job_id not in jobs or jobs[job_id]['status'] != 'completed':
        return jsonify({'error': 'File not ready for download'}), 404

    job_result_dir = jobs[job_id]['result']['job_result_dir']
    job_result = jobs[job_id]['result']

    # 根据请求参数确定下载类型
    file_type = request.args.get('type', 'markdown')

    try:
        if file_type == 'word':
            # 下载Word文档
            word_filename = job_result['word_filename']
            if os.path.exists(os.path.join(job_result_dir, word_filename)):
                return send_from_directory(job_result_dir, word_filename, as_attachment=True)
            else:
                return jsonify({'error': 'Word文档不存在'}), 404

        elif file_type == 'pdf':
            # 下载带文本的PDF
            if 'text_pdf_filename' in job_result:
                text_pdf_filename = job_result['text_pdf_filename']
                if os.path.exists(os.path.join(job_result_dir, text_pdf_filename)):
                    return send_from_directory(job_result_dir, text_pdf_filename, as_attachment=True)
            return jsonify({'error': '带文本PDF不存在'}), 404

        elif file_type == 'zip':
            # 下载完整压缩包
            if 'complete_zip_filename' in job_result:
                complete_zip_filename = job_result['complete_zip_filename']
                if os.path.exists(os.path.join(job_result_dir, complete_zip_filename)):
                    return send_from_directory(job_result_dir, complete_zip_filename, as_attachment=True)
            return jsonify({'error': '完整压缩包不存在'}), 404

        else:  # markdown (默认)
            # 下载Markdown文件
            filename = job_result['filename']

            # 检查是否有图片，如果有则创建包含图片的zip文件
            images_dir = os.path.join(job_result_dir, 'images')
            if os.path.exists(images_dir) and os.listdir(images_dir):
                import zipfile
                zip_filename = f"{os.path.splitext(filename)[0]}_with_images.zip"
                zip_path = os.path.join(job_result_dir, zip_filename)

                # 如果zip文件不存在，创建它
                if not os.path.exists(zip_path):
                    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                        # 添加markdown文件
                        md_path = os.path.join(job_result_dir, filename)
                        zipf.write(md_path, arcname=filename)

                        # 添加所有图片
                        for img_file in os.listdir(images_dir):
                            img_path = os.path.join(images_dir, img_file)
                            zipf.write(img_path, arcname=os.path.join('images', img_file))

                return send_from_directory(job_result_dir, zip_filename, as_attachment=True)

            # 如果没有图片，直接返回markdown文件
            if os.path.exists(os.path.join(job_result_dir, filename)):
                return send_from_directory(job_result_dir, filename, as_attachment=True)
            else:
                return jsonify({'error': 'Markdown文件不存在'}), 404

    except Exception as e:
        logger.error(f"下载文件失败: {str(e)}")
        return jsonify({'error': f'下载失败: {str(e)}'}), 500

@app.route('/batch_size', methods=['POST'])
def set_batch_size():
    """设置批处理大小"""
    try:
        data = request.get_json()
        if 'batch_size' in data:
            batch_size = int(data['batch_size'])
            if batch_size < 1:
                batch_size = 1
            elif batch_size > 10:
                batch_size = 10

            app.config['BATCH_SIZE'] = batch_size
            logger.info(f"批处理大小已设置为: {batch_size}")

            # 如果有排队的作业，尝试处理
            process_next_job()

            return jsonify({'success': True, 'batch_size': batch_size})
    except Exception as e:
        logger.error(f"设置批处理大小失败: {str(e)}")
        return jsonify({'error': str(e)}), 400

    return jsonify({'error': 'Invalid request'}), 400

@app.route('/optimization_settings', methods=['POST'])
def update_optimization_settings():
    """更新优化设置"""
    try:
        data = request.get_json()
        
        # 更新多进程设置
        if 'enable_multiprocessing' in data:
            app.config['ENABLE_MULTIPROCESSING'] = bool(data['enable_multiprocessing'])
            logger.info(f"多进程处理已{'启用' if app.config['ENABLE_MULTIPROCESSING'] else '禁用'}")
        
        # 更新混合精度设置
        if 'enable_mixed_precision' in data:
            app.config['ENABLE_MIXED_PRECISION'] = bool(data['enable_mixed_precision'])
            logger.info(f"混合精度已{'启用' if app.config['ENABLE_MIXED_PRECISION'] else '禁用'}")
        
        # 更新内存优化设置
        if 'optimize_memory' in data:
            app.config['OPTIMIZE_MEMORY'] = bool(data['optimize_memory'])
            logger.info(f"内存优化已{'启用' if app.config['OPTIMIZE_MEMORY'] else '禁用'}")
        
        # 更新模型预加载设置
        if 'preload_models' in data:
            app.config['PRELOAD_MODELS'] = bool(data['preload_models'])
            logger.info(f"模型预加载已{'启用' if app.config['PRELOAD_MODELS'] else '禁用'}")
            
            # 如果禁用预加载，清空模型缓存
            if not app.config['PRELOAD_MODELS']:
                with model_cache_lock:
                    model_cache.clear()
                    logger.info("模型缓存已清空")
        
        # 更新进程池大小
        if 'process_pool_size' in data:
            pool_size = int(data['process_pool_size'])
            if 1 <= pool_size <= 16:
                app.config['PROCESS_POOL_SIZE'] = pool_size
                logger.info(f"进程池大小已设置为: {pool_size}")
        
        # 更新GPU批处理大小
        if 'gpu_batch_size' in data:
            gpu_batch_size = int(data['gpu_batch_size'])
            if 1 <= gpu_batch_size <= 32:
                app.config['GPU_BATCH_SIZE'] = gpu_batch_size
                logger.info(f"GPU批处理大小已设置为: {gpu_batch_size}")
        
        # 更新CPU批处理大小
        if 'cpu_batch_size' in data:
            cpu_batch_size = int(data['cpu_batch_size'])
            if 1 <= cpu_batch_size <= 16:
                app.config['CPU_BATCH_SIZE'] = cpu_batch_size
                logger.info(f"CPU批处理大小已设置为: {cpu_batch_size}")
        
        return jsonify({
            'success': True,
            'settings': {
                'enable_multiprocessing': app.config['ENABLE_MULTIPROCESSING'],
                'enable_mixed_precision': app.config['ENABLE_MIXED_PRECISION'],
                'optimize_memory': app.config['OPTIMIZE_MEMORY'],
                'preload_models': app.config['PRELOAD_MODELS'],
                'process_pool_size': app.config['PROCESS_POOL_SIZE'],
                'gpu_batch_size': app.config['GPU_BATCH_SIZE'],
                'cpu_batch_size': app.config['CPU_BATCH_SIZE']
            }
        })
        
    except Exception as e:
        logger.error(f"更新优化设置失败: {str(e)}")
        return jsonify({'error': str(e)}), 400

@app.route('/clear_model_cache', methods=['POST'])
def clear_model_cache():
    """清空模型缓存"""
    try:
        with model_cache_lock:
            cache_size = len(model_cache)
            model_cache.clear()
            logger.info(f"已清空 {cache_size} 个缓存模型")
        
        return jsonify({
            'success': True,
            'message': f'已清空 {cache_size} 个缓存模型'
        })
        
    except Exception as e:
        logger.error(f"清空模型缓存失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/performance_stats')
def get_performance_stats():
    """获取性能统计"""
    try:
        # 获取实时系统信息
        import psutil
        
        system_stats = {
            'cpu_percent': psutil.cpu_percent(interval=0.1),
            'memory_percent': psutil.virtual_memory().percent,
            'memory_available_gb': psutil.virtual_memory().available / 1024**3
        }
        
        # 获取GPU信息（如果可用）
        gpu_stats = {}
        try:
            import torch
            if torch.cuda.is_available():
                for i in range(torch.cuda.device_count()):
                    if torch.cuda.is_initialized():
                        gpu_stats[f'gpu_{i}'] = {
                            'memory_allocated_gb': torch.cuda.memory_allocated(i) / 1024**3,
                            'memory_reserved_gb': torch.cuda.memory_reserved(i) / 1024**3,
                            'memory_total_gb': torch.cuda.get_device_properties(i).total_memory / 1024**3
                        }
        except ImportError:
            pass
        
        return jsonify({
            'performance_stats': performance_stats,
            'system_stats': system_stats,
            'gpu_stats': gpu_stats,
            'model_cache_size': len(model_cache),
            'active_jobs': active_jobs,
            'queued_jobs': len(job_queue)
        })
        
    except Exception as e:
        logger.error(f"获取性能统计失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/auto_optimize', methods=['POST'])
def auto_optimize():
    """自动优化配置"""
    try:
        optimal_config, recommendations = auto_configure_optimization()
        
        # 应用优化配置
        app.config['ENABLE_MULTIPROCESSING'] = optimal_config.get('enable_multiprocessing', False)
        app.config['ENABLE_MIXED_PRECISION'] = optimal_config.get('enable_mixed_precision', False)
        app.config['OPTIMIZE_MEMORY'] = optimal_config.get('optimize_memory', True)
        app.config['PRELOAD_MODELS'] = optimal_config.get('preload_models', False)
        app.config['PROCESS_POOL_SIZE'] = optimal_config.get('process_pool_size', 2)
        app.config['GPU_BATCH_SIZE'] = optimal_config.get('gpu_batch_size', 4)
        app.config['CPU_BATCH_SIZE'] = optimal_config.get('cpu_batch_size', 2)
        app.config['MAX_WORKERS'] = optimal_config.get('max_workers', 4)
        
        # 如果有GPU，自动启用GPU
        if optimal_config.get('device') == 'cuda':
            app.config['USE_GPU'] = True
        
        logger.info("自动优化配置已应用")
        
        return jsonify({
            'success': True,
            'applied_config': optimal_config,
            'recommendations': recommendations,
            'message': '自动优化配置已应用'
        })
        
    except Exception as e:
        logger.error(f"自动优化失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/system_performance_report')
def system_performance_report():
    """获取系统性能报告"""
    try:
        report = get_system_performance_report()
        return jsonify(report)
    except Exception as e:
        logger.error(f"获取系统性能报告失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/preset_configs')
def get_preset_configs():
    """获取预设配置列表"""
    try:
        presets = list_preset_configs()
        return jsonify({'presets': presets})
    except Exception as e:
        logger.error(f"获取预设配置失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/apply_preset/<preset_name>', methods=['POST'])
def apply_preset_config(preset_name):
    """应用预设配置"""
    try:
        preset = get_preset_config(preset_name)
        if not preset:
            return jsonify({'error': '预设配置不存在'}), 404
        
        config = preset['config']
        
        # 应用预设配置
        if 'enable_multiprocessing' in config:
            app.config['ENABLE_MULTIPROCESSING'] = config['enable_multiprocessing']
        if 'enable_mixed_precision' in config:
            app.config['ENABLE_MIXED_PRECISION'] = config['enable_mixed_precision']
        if 'optimize_memory' in config:
            app.config['OPTIMIZE_MEMORY'] = config['optimize_memory']
        if 'preload_models' in config:
            app.config['PRELOAD_MODELS'] = config['preload_models']
        if 'batch_size' in config:
            app.config['BATCH_SIZE'] = config['batch_size']
        if 'process_pool_size' in config:
            app.config['PROCESS_POOL_SIZE'] = config['process_pool_size']
        
        # 应用OCR配置
        if 'ocr_level' in config:
            current_config['ocr_level'] = config['ocr_level']
        if 'extract_table_format' in config:
            current_config['extract_table_format'] = config['extract_table_format']
        
        logger.info(f"已应用预设配置: {preset_name}")
        
        return jsonify({
            'success': True,
            'preset_name': preset_name,
            'preset_info': preset,
            'applied_config': config,
            'message': f'已应用预设配置: {preset["name"]}'
        })
        
    except Exception as e:
        logger.error(f"应用预设配置失败: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/run_mode', methods=['GET', 'POST'])
def run_mode():
    """获取或设置运行模式"""
    if request.method == 'GET':
        # 获取当前运行模式
        current_mode = 'compatible'  # 默认兼容模式
        if app.config.get('ENABLE_MULTIPROCESSING') and app.config.get('USE_GPU'):
            current_mode = 'performance'
        elif app.config.get('USE_GPU') or app.config.get('ENABLE_MULTIPROCESSING'):
            current_mode = 'balanced'
        
        return jsonify({
            'current_mode': current_mode,
            'modes': {
                'compatible': {
                    'name': '兼容模式',
                    'description': '最高稳定性，适合所有环境'
                },
                'balanced': {
                    'name': '平衡模式',
                    'description': '平衡性能和稳定性'
                },
                'performance': {
                    'name': '性能模式',
                    'description': '最大化处理速度'
                }
            }
        })
    
    elif request.method == 'POST':
        # 设置运行模式
        data = request.get_json()
        mode = data.get('mode', 'compatible')
        
        if mode == 'compatible':
            # 兼容模式：禁用所有高级功能
            app.config['USE_GPU'] = False
            app.config['ENABLE_MULTIPROCESSING'] = False
            app.config['ENABLE_MIXED_PRECISION'] = False
            app.config['PRELOAD_MODELS'] = False
            app.config['BATCH_SIZE'] = 3
            
        elif mode == 'balanced':
            # 平衡模式：启用部分优化
            app.config['USE_GPU'] = True
            app.config['ENABLE_MULTIPROCESSING'] = False
            app.config['ENABLE_MIXED_PRECISION'] = False
            app.config['PRELOAD_MODELS'] = True
            app.config['BATCH_SIZE'] = 5
            
        elif mode == 'performance':
            # 性能模式：启用所有优化
            app.config['USE_GPU'] = True
            app.config['ENABLE_MULTIPROCESSING'] = True
            app.config['ENABLE_MIXED_PRECISION'] = True
            app.config['PRELOAD_MODELS'] = True
            app.config['BATCH_SIZE'] = 8
            
        else:
            return jsonify({'error': '无效的运行模式'}), 400
        
        logger.info(f"运行模式已切换到: {mode}")
        
        return jsonify({
            'success': True,
            'mode': mode,
            'message': f'已切换到{mode}模式'
        })

# 应用启动时加载作业状态
load_jobs()

# 启动时处理队列中的作业
def start_processing_queue():
    logger.info("启动队列处理")
    process_next_job()

# 启动队列处理线程
threading.Thread(target=start_processing_queue).start()

if __name__ == '__main__':
    logger.info(f"Starting application with {'GPU' if app.config['USE_GPU'] else 'CPU'} acceleration")
    logger.info(f"Model directory: {app.config['MODEL_DIR']}")
    app.run(debug=True)
